"""Create a file containing the verification and validation of Expressions report from e2e tests HTML report file.

This file is then imported into a Confluence page to be filled out by the user.

If saved as DOCX:
    This can be done by creating a new page and clicking the three dots in the top right corner, then selecting "Templates & import doc".
If saved as QMD:
    This can be done by using Quarto for automated publishing to Confluence.

Input arguments:
    --file-path: The path to the HTML report file.
    --output-path: The path of the output file.
    --format: The format of the output file. Default is "DOCX", the other option is "QMD".
    --genialis-base-version: The Genialis base version used in the tested deploy.
"""

import argparse
import logging
import os
import pwd
import re
import time
from datetime import datetime
from io import StringIO
from pathlib import Path

import pandas as pd
import pypandoc
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import RGBColor
from selenium import webdriver
from selenium.webdriver.chrome.options import Options

logging.basicConfig(level=logging.INFO)

DEPLOY_MAPPINGS = {
    "https://qa.dictyexpress.org/": "QA",
    "https://app.dictyexpress.org/": "APP",
}

TEMPLATE_VERSION = "1.1.0"

parser = argparse.ArgumentParser(
    description="Convert HTML e2e report to QMD or DOCX file."
)

parser.add_argument(
    "-i",
    "--file-path",
    type=str,
    required=True,
    help="The path to the input HTML report file.",
)
parser.add_argument(
    "-o",
    "--output-path",
    type=str,
    required=True,
    help="The path in which to save the output file. File name will be the same as the input file name.",
)
parser.add_argument(
    "-f",
    "--format",
    type=str,
    default="DOCX",
    help="The format of the output file.",
    choices=["QMD", "DOCX"],
)
parser.add_argument(
    "-v",
    "--genialis-base-version",
    type=str,
    help="Genialis base version used in the tested deploy.",
    required=True,
)


def make_bold(cell):
    """Make the text in the cell bold."""
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True


def get_number(summary_div, class_name):
    text = summary_div.find("span", class_=class_name).text
    return int(re.search(r"\d+", text).group())


args = parser.parse_args()

if __name__ == "__main__":
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")

    driver = webdriver.Chrome(options=chrome_options)

    html_file_path = Path(args.file_path).resolve()
    out_file_path = Path(args.output_path).resolve()
    out_file_path.mkdir(parents=True, exist_ok=True)

    if not html_file_path.is_file():
        logging.error(f"File {html_file_path} not found")
        raise FileNotFoundError

    if out_file_path.suffix:
        raise ValueError("Output path should be a directory.")
    else:
        out_file_path = out_file_path / html_file_path.stem

    driver.get(f"file://{html_file_path.as_posix()}")

    # Wait for the page to fully load
    time.sleep(2)

    html = driver.page_source
    soup = BeautifulSoup(html, "html.parser")
    e2e_runtime_str = soup.find("p")
    summary_div = soup.find("div", class_="filters")

    passed = get_number(summary_div, "passed")
    failed = get_number(summary_div, "failed")
    skipped = get_number(summary_div, "skipped")
    unexpected_passed = get_number(summary_div, "xpassed")
    expected_failed = get_number(summary_div, "xfailed")
    error = get_number(summary_div, "error")

    e2e_tables = pd.read_html(StringIO(html))

    doc = Document()

    # Summary and deployment notes
    doc.add_heading("Summary and deployment notes", 1)

    doc.add_paragraph(
        f"Test statistics: Passed- {passed}, Failed- {failed}, Skipped- {skipped}, Unexpected passed- {unexpected_passed}, Expected failed- {expected_failed}, Error- {error}"
    )

    # General information
    doc.add_heading("General information", level=1)

    uid = os.getuid()
    user_info = pwd.getpwuid(uid)
    full_name = user_info.pw_gecos

    general_table = doc.add_table(rows=2, cols=2)
    general_table.style = "Table Grid"
    general_table.rows[0].cells[0].text = "Tester"
    general_table.rows[0].cells[1].text = full_name
    general_table.rows[1].cells[0].text = "Template version"
    general_table.rows[1].cells[1].text = TEMPLATE_VERSION

    info_texts = [
        "Deploy date",
        "Testing date",
    ]
    info_table = doc.add_table(rows=len(info_texts), cols=2)
    info_table.style = "Table Grid"

    for i, row in enumerate(info_table.rows):
        cells = row.cells
        cells[0].text = info_texts[i]
        # Set deploy and testing date to current date (easier to change post generation)
        if info_texts[i] == "Deploy date" or info_texts[i] == "Testing date":
            today_date = datetime.today()
            date_string = today_date.strftime("%Y-%m-%d")
            cells[1].text = date_string

    doc.add_heading("e2e test report", level=1)

    # Example paragraph: Report generated on 09-Jul-2024 at 12:07:56 by pytest-html v4.1.1
    doc.add_paragraph(re.sub(r"\s+", " ", e2e_runtime_str.get_text()))

    # First table is the e2e environment information
    env_info = e2e_tables[0]

    env_info_table = doc.add_table(
        rows=len(env_info.values), cols=len(env_info.columns)
    )

    # Header row
    env_info_table.style = "Table Grid"
    hdr_cells = env_info_table.rows[0].cells
    hdr_cells[0].text = "Module"
    hdr_cells[1].text = "Version"
    make_bold(hdr_cells[0])
    make_bold(hdr_cells[1])

    # Subsequent rows
    re_pattern = re.compile(r"([\w-]+): ([\d.]+)")
    for i, row in enumerate(env_info.values):
        if row[0] == "Packages" or row[0] == "Plugins":
            value = row[1]
            modules_list = re_pattern.findall(str(value))
            modules_paragraph = "\n".join(
                f"{module}: {version}" for module, version in modules_list
            )
            row[1] = modules_paragraph
        elif row[0] == "Base URL":
            base_url = row[1]
            deploy = DEPLOY_MAPPINGS[base_url]

        cells = env_info_table.rows[i].cells
        for j, value in enumerate(row):
            if str(value) == "nan":
                cells[j].text = ""
            else:
                cells[j].text = str(value)

    # Add a new row for the user to attach the origin HTML report
    new_row = env_info_table.add_row().cells
    cells = env_info_table.rows[-1].cells
    cells[0].text = "e2e HTML Report"

    # Second table are the actual test results
    results = e2e_tables[1]
    # Remove rows with "<  >" in them, because they include log information
    # These rows can be accessed by opening the HTML report
    table = results[~results.map(lambda x: "<  >" in str(x)).any(axis=1)]

    # Group the tests into test groups
    table = table.copy()
    table["group"] = table["Test"].apply(lambda x: x.split("::")[0])
    grouped = table.groupby("group")

    # Every test group should contain a separate table
    for name, group in grouped:
        # Do not include passed tests in the report
        group = group[group["Result"] != "Passed"]

        heading = doc.add_heading(name, level=2)
        if group.empty:
            checkmark = heading.add_run(" \u2714")
            checkmark.font.color.rgb = RGBColor(0, 128, 0)
            continue
        else:
            cross = heading.add_run(" \u2716")
            cross.font.color.rgb = RGBColor(255, 0, 0)
        group = group.drop(columns=["group"])

        tests_table = doc.add_table(rows=len(group.values) + 1, cols=len(group.columns))
        tests_table.style = "Table Grid"

        # Header row with column names: Result, Test, Duration, Ticket Link/Comments
        hdr_cells = tests_table.rows[0].cells
        for j, column in enumerate(group.columns):
            if str(column) == "Links":
                hdr_cells[j].text = "Ticket Link/Comments"
            else:
                hdr_cells[j].text = str(column)

            make_bold(hdr_cells[j])

        # Actual test information rows
        for i, row in enumerate(group.values, start=1):
            cells = tests_table.rows[i].cells
            for j, value in enumerate(row):
                if j == 0:
                    cell = cells[j]
                    run = cell.paragraphs[0].add_run(str(value))
                    # Color for Skipped and Expected failed should be different than for Failed and Error
                    if str(value).startswith("Skipped") or str(value).startswith(
                        "Expected"
                    ):
                        run.font.color.rgb = RGBColor(255, 165, 0)
                    else:
                        run.font.color.rgb = RGBColor(255, 0, 0)
                elif j == 1:
                    text = str(value).split("::")[1]
                    # Replace chromium-False with Public user and chromium-True with Logged in user
                    if "[chromium-False]" in text:
                        text = text.replace("[chromium-False]", " [Public user]")
                    elif "[chromium-True]" in text:
                        text = text.replace("[chromium-True]", " [Logged in user]")
                    cells[j].text = text
                else:
                    if str(value) == "nan":
                        cells[j].text = ""
                    else:
                        cells[j].text = str(value)

    docx_path = out_file_path.with_suffix(".docx")
    doc.save(docx_path)
    if args.format == "QMD":
        qmd_path = out_file_path.with_suffix(".qmd")
        title = title = (
            "---\n"
            f"title: {date_string} | {args.genialis_base_version} | DictyExpress {deploy} | E2E\n"
            "---\n\n"
        )
        pypandoc.convert_file(docx_path, "md", outputfile=qmd_path)
        with open(qmd_path, "r+") as file:
            content = file.read()
            file.seek(0, 0)
            file.write(title + content)

        docx_path.unlink()
        logging.info(f"File saved to {qmd_path}")
    else:
        logging.info(f"File saved to {docx_path}")

    driver.quit()
